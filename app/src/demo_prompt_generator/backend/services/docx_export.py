"""Convert walkthrough markdown to a styled Word document."""

from __future__ import annotations

import io
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn


def _add_horizontal_rule(doc: Document):
    """Add a thin horizontal line as a paragraph border."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(
        qn("w:bottom"),
        {
            qn("w:val"): "single",
            qn("w:sz"): "6",
            qn("w:space"): "1",
            qn("w:color"): "CCCCCC",
        },
    )
    pBdr.append(bottom)
    pPr.append(pBdr)


def _style_navigation_cue(paragraph, cue_type: str, cue_text: str):
    """Style a navigation cue as a colored inline run."""
    run = paragraph.add_run()
    if cue_type == "navigate":
        run.text = f"  ▸ Navigate to {cue_text}"
        run.font.color.rgb = RGBColor(0x33, 0x7A, 0xB7)  # blue
    else:
        run.text = f"  ◉ Show {cue_text}"
        run.font.color.rgb = RGBColor(0xD4, 0x8A, 0x0C)  # amber
    run.font.size = Pt(10)
    run.bold = True


def walkthrough_md_to_docx(markdown: str, demo_name: str = "Demo") -> bytes:
    """Convert walkthrough markdown content to a .docx file and return bytes."""
    doc = Document()

    # -- Page setup --
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # -- Style defaults --
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    for level in range(1, 4):
        hstyle = doc.styles[f"Heading {level}"]
        hstyle.font.name = "Calibri"
        hstyle.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    # -- Parse and render --
    lines = markdown.split("\n")
    i = 0
    in_code_block = False
    code_buffer: list[str] = []

    while i < len(lines):
        line = lines[i]

        # Code block toggle
        if line.strip().startswith("```"):
            if in_code_block:
                # End code block — render as styled paragraph
                code_text = "\n".join(code_buffer)
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(code_text)
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x6B, 0x4F, 0xBB)  # violet
                in_code_block = False
                code_buffer = []
            else:
                in_code_block = True
                code_buffer = []
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # H1
        if line.startswith("# ") and not line.startswith("## "):
            title = line[2:].strip()
            h = doc.add_heading(title, level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Subtitle with demo name
            sub = doc.add_paragraph()
            run = sub.add_run(f"Demo Package: {demo_name}")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            run.italic = True
            _add_horizontal_rule(doc)
            i += 1
            continue

        # H2
        if line.startswith("## "):
            _add_horizontal_rule(doc)
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue

        # H3
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue

        # Navigation cues
        nav_match = re.search(r"\[Navigate to (.+?)\]", line)
        show_match = re.search(r"\[Show (.+?)\]", line)
        if nav_match:
            p = doc.add_paragraph()
            _style_navigation_cue(p, "navigate", nav_match.group(1))
            i += 1
            continue
        if show_match:
            p = doc.add_paragraph()
            _style_navigation_cue(p, "show", show_match.group(1))
            i += 1
            continue

        # Bullet points
        bullet_match = re.match(r"^(\s*)[-*]\s+(.+)", line)
        if bullet_match:
            text = bullet_match.group(2)
            p = doc.add_paragraph(style="List Bullet")
            # Handle bold fragments
            parts = re.split(r"(\*\*.+?\*\*)", text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
            i += 1
            continue

        # Numbered list
        num_match = re.match(r"^\d+\.\s+(.+)", line)
        if num_match:
            text = num_match.group(1)
            p = doc.add_paragraph(style="List Number")
            parts = re.split(r"(\*\*.+?\*\*)", text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
            i += 1
            continue

        # Regular paragraph (skip empty lines)
        if line.strip():
            p = doc.add_paragraph()
            # Handle inline bold and code
            parts = re.split(r"(\*\*.+?\*\*|`.+?`)", line)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith("`") and part.endswith("`"):
                    run = p.add_run(part[1:-1])
                    run.font.name = "Consolas"
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0x6B, 0x4F, 0xBB)
                else:
                    p.add_run(part)

        i += 1

    # Write to bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
